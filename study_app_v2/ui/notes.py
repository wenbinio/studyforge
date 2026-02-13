"""notes.py — Notes Manager tab for StudyForge v2.
Import .txt, .md, .pdf, .docx files; view, tag, search, and manage notes.
Includes rich markdown editing, preview, navigator, focus mode,
document export, find & replace, table insertion, and document stats.
"""

import customtkinter as ctk
import threading, os, re
from tkinter import filedialog
from ui.styles import COLORS, FONTS, PAD
import database as db


def extract_text(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    if ext in (".txt", ".md"):
        with open(filepath, "r", encoding="utf-8", errors="replace") as f: return f.read()
    elif ext == ".pdf":
        try:
            import fitz
            with fitz.open(filepath) as doc:
                t = "\n".join(p.get_text() for p in doc)
            return t.strip()
        except ImportError: return "[PyMuPDF not installed]"
    elif ext == ".docx":
        try:
            from docx import Document
            return "\n".join(p.text for p in Document(filepath).paragraphs)
        except ImportError: return "[python-docx not installed]"
    return f"[Unsupported: {ext}]"


class NotesTab(ctk.CTkFrame):
    def __init__(self, parent, app_ref):
        super().__init__(parent, fg_color="transparent")
        self.app = app_ref
        self.sel_id = None
        self.preview_visible = False
        self.nav_visible = False
        self.build_ui()

    def build_ui(self):
        header = ctk.CTkFrame(self, fg_color="transparent")
        header.pack(fill="x", padx=PAD["page"], pady=(PAD["page"], 5))
        ctk.CTkLabel(header, text="📝 Notes", font=FONTS["heading"],
            text_color=COLORS["text_primary"]).pack(side="left")

        br = ctk.CTkFrame(header, fg_color="transparent")
        br.pack(side="right")

        self.focus_btn = ctk.CTkButton(br, text="🖥️ Focus Mode", width=120, height=34,
            font=FONTS["body"], fg_color=COLORS["bg_card"], hover_color=COLORS["accent_hover"],
            corner_radius=8, command=self._toggle_focus_mode)
        self.focus_btn.pack(side="left", padx=4)

        ctk.CTkButton(br, text="➕ New Note", width=120, height=34, font=FONTS["body"],
            fg_color=COLORS["success"], hover_color="#00d2a0", corner_radius=8,
            command=self._new_note).pack(side="left", padx=4)
        ctk.CTkButton(br, text="📂 Import File", width=120, height=34, font=FONTS["body"],
            fg_color=COLORS["accent"], hover_color=COLORS["accent_hover"], corner_radius=8,
            command=self.import_file).pack(side="left", padx=4)
        ctk.CTkButton(br, text="📋 Paste Note", width=120, height=34, font=FONTS["body"],
            fg_color=COLORS["success"], hover_color="#00d2a0", corner_radius=8,
            command=self.paste_dlg).pack(side="left", padx=4)
        self.export_btn = ctk.CTkButton(br, text="📤 Export", width=100, height=34,
            font=FONTS["body"], fg_color=COLORS["bg_card"], hover_color=COLORS["accent_hover"],
            corner_radius=8, command=self._export_note, state="disabled")
        self.export_btn.pack(side="left", padx=4)

        sf = ctk.CTkFrame(self, fg_color="transparent")
        sf.pack(fill="x", padx=PAD["page"], pady=(5,8))
        self.search_var = ctk.StringVar()
        self.search_var.trace_add("write", lambda *_: self.refresh())
        ctk.CTkEntry(sf, textvariable=self.search_var, placeholder_text="🔍 Search notes...",
            fg_color=COLORS["bg_input"], text_color=COLORS["text_primary"], font=FONTS["body"],
            border_color=COLORS["border"], corner_radius=8, height=36).pack(fill="x")

        split = ctk.CTkFrame(self, fg_color="transparent")
        split.pack(fill="both", expand=True, padx=PAD["page"], pady=(0, PAD["page"]))
        split.grid_columnconfigure(0, weight=1, minsize=280)
        split.grid_columnconfigure(1, weight=3)
        split.grid_rowconfigure(0, weight=1)

        self.list_f = ctk.CTkScrollableFrame(split, fg_color=COLORS["bg_card"], corner_radius=12, width=280)
        self.list_f.grid(row=0, column=0, sticky="nsew", padx=(0,8))
        self.viewer = ctk.CTkFrame(split, fg_color=COLORS["bg_card"], corner_radius=12)
        self.viewer.grid(row=0, column=1, sticky="nsew")

        self.refresh()
        self._empty_viewer()

    def _empty_viewer(self):
        for w in self.viewer.winfo_children(): w.destroy()
        try:
            self.export_btn.configure(state="disabled")
        except Exception:
            pass
        ctk.CTkLabel(self.viewer, text="📄", font=("Segoe UI", 40),
            text_color=COLORS["text_muted"]).pack(expand=True)
        ctk.CTkLabel(self.viewer, text="Select a note, or click ➕ New Note to create one.",
            font=FONTS["body"], text_color=COLORS["text_muted"]).pack(pady=(0,40))

    def refresh(self):
        for w in self.list_f.winfo_children(): w.destroy()
        q = self.search_var.get().strip()
        notes = db.search_notes(q) if q else db.get_all_notes()
        if not notes:
            ctk.CTkLabel(self.list_f, text="No notes yet.", font=FONTS["body"],
                text_color=COLORS["text_muted"]).pack(pady=40)
            return
        for n in notes:
            sel = n["id"] == self.sel_id
            bg = COLORS["accent"] if sel else COLORS["bg_secondary"]
            item = ctk.CTkFrame(self.list_f, fg_color=bg, corner_radius=8, cursor="hand2")
            item.pack(fill="x", padx=6, pady=3)
            t = n["title"][:40] + ("..." if len(n["title"]) > 40 else "")
            ctk.CTkLabel(item, text=t, font=FONTS["body_bold"],
                text_color=COLORS["text_primary"], anchor="w").pack(padx=10, pady=(8,0), anchor="w")
            meta = f"{n.get('tags','') or 'No tags'}  ·  {n['created_at'][:10]}"
            ctk.CTkLabel(item, text=meta, font=FONTS["small"],
                text_color=COLORS["text_muted"], anchor="w").pack(padx=10, pady=(0,6), anchor="w")
            for w in [item] + item.winfo_children():
                w.bind("<Button-1>", lambda e, nid=n["id"]: self.view_note(nid))

    def view_note(self, nid):
        self.sel_id = nid; self.refresh()
        self.export_btn.configure(state="normal")
        note = db.get_note(nid)
        if not note: return
        for w in self.viewer.winfo_children(): w.destroy()

        tb = ctk.CTkFrame(self.viewer, fg_color="transparent")
        tb.pack(fill="x", padx=PAD["section"], pady=(PAD["section"],5))
        self.title_e = ctk.CTkEntry(tb, font=FONTS["subheading"], fg_color=COLORS["bg_input"],
            text_color=COLORS["text_primary"], border_color=COLORS["border"], corner_radius=8)
        self.title_e.insert(0, note["title"])
        self.title_e.pack(side="left", fill="x", expand=True, padx=(0,8))
        ctk.CTkButton(tb, text="💾", width=36, height=32, font=FONTS["small"],
            fg_color=COLORS["success"], corner_radius=6,
            command=lambda: self._save(nid)).pack(side="left", padx=2)
        ctk.CTkButton(tb, text="🗑️", width=36, height=32, font=FONTS["small"],
            fg_color=COLORS["danger"], corner_radius=6,
            command=lambda: self._del(nid)).pack(side="left", padx=2)

        self.preview_btn = ctk.CTkButton(tb, text="👁️", width=36, height=32,
            font=FONTS["small"],
            fg_color=COLORS["accent"] if self.preview_visible else COLORS["bg_card"],
            hover_color=COLORS["accent_hover"], corner_radius=6,
            command=self._toggle_preview)
        self.preview_btn.pack(side="left", padx=2)

        self.nav_btn = ctk.CTkButton(tb, text="Nav", width=42, height=32,
            font=FONTS["small"],
            fg_color=COLORS["accent"] if self.nav_visible else COLORS["bg_card"],
            hover_color=COLORS["accent_hover"], corner_radius=6,
            command=self._toggle_navigator)
        self.nav_btn.pack(side="left", padx=2)

        ctk.CTkButton(tb, text="🔎", width=36, height=32, font=FONTS["small"],
            fg_color=COLORS["bg_card"], hover_color=COLORS["accent_hover"],
            corner_radius=6, command=self._show_find_replace).pack(side="left", padx=2)

        tgf = ctk.CTkFrame(self.viewer, fg_color="transparent")
        tgf.pack(fill="x", padx=PAD["section"], pady=(0,5))
        ctk.CTkLabel(tgf, text="Tags:", font=FONTS["small"],
            text_color=COLORS["text_muted"]).pack(side="left")
        self.tags_e = ctk.CTkEntry(tgf, font=FONTS["small"], fg_color=COLORS["bg_input"],
            text_color=COLORS["text_primary"], border_color=COLORS["border"], corner_radius=6, height=28)
        self.tags_e.insert(0, note.get("tags",""))
        self.tags_e.pack(side="left", fill="x", expand=True, padx=6)

        ctk.CTkLabel(tgf, text="Font:", font=FONTS["small"],
            text_color=COLORS["text_muted"]).pack(side="left", padx=(12, 2))
        self._font_size_var = ctk.StringVar(value="13")
        ctk.CTkOptionMenu(tgf, variable=self._font_size_var,
            values=["10", "11", "12", "13", "14", "16", "18", "20"],
            width=60, height=28, font=FONTS["small"],
            fg_color=COLORS["bg_input"], button_color=COLORS["accent"],
            button_hover_color=COLORS["accent_hover"],
            dropdown_fg_color=COLORS["bg_card"],
            dropdown_hover_color=COLORS["accent_hover"],
            command=self._change_font_size).pack(side="left", padx=2)

        ar = ctk.CTkFrame(self.viewer, fg_color="transparent")
        ar.pack(fill="x", padx=PAD["section"], pady=(0,5))
        ctk.CTkButton(ar, text="🤖 Generate Cards", width=140, height=30, font=FONTS["small"],
            fg_color=COLORS["warning"], corner_radius=6,
            command=lambda: self._gen_cards(note)).pack(side="left", padx=2)
        ctk.CTkButton(ar, text="📋 Summarize", width=110, height=30, font=FONTS["small"],
            fg_color=COLORS["accent"], corner_radius=6,
            command=lambda: self._summarize(note)).pack(side="left", padx=2)
        ctk.CTkButton(ar, text="❓ Ask Question", width=120, height=30, font=FONTS["small"],
            fg_color=COLORS["accent_light"], corner_radius=6,
            command=lambda: self._ask(note)).pack(side="left", padx=2)

        # Formatting toolbar
        toolbar = ctk.CTkFrame(self.viewer, fg_color=COLORS["bg_secondary"], corner_radius=8, height=36)
        toolbar.pack(fill="x", padx=PAD["section"], pady=(0, 5))

        fmt_buttons = [
            ("H1", lambda: self._insert_prefix("# ")),
            ("H2", lambda: self._insert_prefix("## ")),
            ("H3", lambda: self._insert_prefix("### ")),
            ("H4", lambda: self._insert_prefix("#### ")),
            ("|", None),
            ("B", lambda: self._wrap_selection("**")),
            ("I", lambda: self._wrap_selection("*")),
            ("~~", lambda: self._wrap_selection("~~")),
            ("|", None),
            ("• List", lambda: self._insert_prefix("- ")),
            ("1. List", lambda: self._insert_prefix("1. ")),
            ("☑ Task", lambda: self._insert_prefix("- [ ] ")),
            ("|", None),
            ("`Code`", lambda: self._wrap_selection("`")),
            ("```", self._insert_code_block),
            ("> Quote", lambda: self._insert_prefix("> ")),
            ("---", lambda: self._insert_line("\n---\n")),
            ("|", None),
            ("🔗", self._insert_link),
            ("🖼️", self._insert_image),
            ("📊", self._insert_table),
        ]

        for text, cmd in fmt_buttons:
            if cmd is None:
                ctk.CTkFrame(toolbar, width=1, height=18, fg_color=COLORS["border"]).pack(
                    side="left", padx=4, pady=7)
            else:
                if text in ("H1", "H2", "H3", "H4") or text == "B":
                    btn_font = ("Segoe UI", 11, "bold")
                elif text == "I":
                    btn_font = ("Segoe UI", 11, "italic")
                else:
                    btn_font = FONTS["small"]
                ctk.CTkButton(toolbar, text=text, width=max(32, len(text) * 9 + 8), height=24,
                    font=btn_font, fg_color="transparent", hover_color=COLORS["bg_hover"],
                    text_color=COLORS["text_secondary"], corner_radius=6,
                    command=cmd).pack(side="left", padx=1, pady=4)

        # Editor area with optional navigator and preview panels
        editor_area = ctk.CTkFrame(self.viewer, fg_color="transparent")
        editor_area.pack(fill="both", expand=True, padx=PAD["section"], pady=(0, 0))
        editor_area.grid_rowconfigure(0, weight=1)
        self._editor_area = editor_area

        self.nav_panel = ctk.CTkScrollableFrame(editor_area, fg_color=COLORS["bg_secondary"],
            width=160, corner_radius=8, border_color=COLORS["border"], border_width=1)

        self.editor = ctk.CTkTextbox(editor_area, fg_color=COLORS["bg_input"],
            text_color=COLORS["text_primary"], font=("Consolas", 13),
            border_color=COLORS["border"], border_width=1, corner_radius=8, wrap="word", undo=True)
        self.editor.insert("1.0", note["content"])

        self.preview_panel = ctk.CTkTextbox(editor_area, fg_color=COLORS["bg_secondary"],
            text_color=COLORS["text_primary"], font=FONTS["body"],
            border_color=COLORS["accent"], border_width=1, corner_radius=8, wrap="word")

        self._update_editor_layout()

        # Document stats bar
        self.stats_bar = ctk.CTkFrame(self.viewer, fg_color=COLORS["bg_secondary"],
            corner_radius=8, height=28)
        self.stats_bar.pack(fill="x", padx=PAD["section"], pady=(4, PAD["section"]))
        self.stats_label = ctk.CTkLabel(self.stats_bar, text="", font=FONTS["small"],
            text_color=COLORS["text_muted"])
        self.stats_label.pack(side="left", padx=10, pady=4)
        self._update_stats()

        # Keyboard shortcuts
        self.editor.bind("<Control-b>", lambda e: (self._wrap_selection("**"), "break"))
        self.editor.bind("<Control-i>", lambda e: (self._wrap_selection("*"), "break"))
        self.editor.bind("<Control-k>", lambda e: (self._wrap_selection("`"), "break"))
        self.editor.bind("<Control-s>", lambda e: (self._save(nid), "break"))
        self.editor.bind("<Control-h>", lambda e: (self._show_find_replace(), "break"))
        self.editor.bind("<Control-B>", lambda e: (self._wrap_selection("**"), "break"))
        self.editor.bind("<Control-I>", lambda e: (self._wrap_selection("*"), "break"))
        self.editor.bind("<Control-K>", lambda e: (self._wrap_selection("`"), "break"))
        self.editor.bind("<Control-S>", lambda e: (self._save(nid), "break"))
        self.editor.bind("<Control-H>", lambda e: (self._show_find_replace(), "break"))
        self.editor.bind("<KeyRelease>", lambda e: self._update_stats())

    # ── Editor layout with navigator/preview ─────────────────────

    def _update_editor_layout(self):
        editor_col = 1 if self.nav_visible else 0
        preview_col = editor_col + 1

        if self.nav_visible:
            self._refresh_navigator()
            self.nav_panel.grid(row=0, column=0, sticky="nsew", padx=(0, 5))
        else:
            self.nav_panel.grid_forget()

        self.editor.grid(row=0, column=editor_col, sticky="nsew")

        if self.preview_visible:
            self._render_preview()
            self.preview_panel.grid(row=0, column=preview_col, sticky="nsew", padx=(5, 0))
        else:
            self.preview_panel.grid_forget()

        for col in range(3):
            self._editor_area.grid_columnconfigure(col, weight=0)
        self._editor_area.grid_columnconfigure(editor_col, weight=1)
        if self.preview_visible:
            self._editor_area.grid_columnconfigure(preview_col, weight=1)

    # ── Formatting helpers ────────────────────────────────────────

    def _insert_prefix(self, prefix):
        try:
            cursor = self.editor.index("insert")
            line = cursor.split(".")[0]
            line_start = f"{line}.0"
            line_text = self.editor.get(line_start, f"{line}.end")
            if line_text.startswith(prefix):
                return
            self.editor.insert(line_start, prefix)
            self.editor.focus_set()
        except Exception:
            pass

    def _wrap_selection(self, wrapper):
        try:
            sel_start = self.editor.index("sel.first")
            sel_end = self.editor.index("sel.last")
            selected = self.editor.get(sel_start, sel_end)
            self.editor.delete(sel_start, sel_end)
            self.editor.insert(sel_start, f"{wrapper}{selected}{wrapper}")
            self.editor.focus_set()
        except Exception:
            cursor = self.editor.index("insert")
            self.editor.insert(cursor, f"{wrapper}{wrapper}")
            new_pos = self.editor.index(f"{cursor} + {len(wrapper)} chars")
            self.editor.mark_set("insert", new_pos)
            self.editor.focus_set()

    def _insert_code_block(self):
        try:
            sel_start = self.editor.index("sel.first")
            sel_end = self.editor.index("sel.last")
            selected = self.editor.get(sel_start, sel_end)
            self.editor.delete(sel_start, sel_end)
            self.editor.insert(sel_start, f"```\n{selected}\n```")
        except Exception:
            cursor = self.editor.index("insert")
            self.editor.insert(cursor, "```\n\n```")
            new_pos = self.editor.index(f"{cursor} + 4 chars")
            self.editor.mark_set("insert", new_pos)
        self.editor.focus_set()

    def _insert_line(self, text):
        cursor = self.editor.index("insert")
        self.editor.insert(cursor, text)
        self.editor.focus_set()

    def _insert_link(self):
        try:
            sel_start = self.editor.index("sel.first")
            sel_end = self.editor.index("sel.last")
            selected = self.editor.get(sel_start, sel_end)
            self.editor.delete(sel_start, sel_end)
            self.editor.insert(sel_start, f"[{selected}](url)")
        except Exception:
            cursor = self.editor.index("insert")
            self.editor.insert(cursor, "[link text](url)")
        self.editor.focus_set()

    def _insert_image(self):
        cursor = self.editor.index("insert")
        self.editor.insert(cursor, "![alt text](image_url)")
        self.editor.focus_set()

    def _insert_table(self):
        win = ctk.CTkToplevel(self)
        win.title("Insert Table"); win.geometry("300x200")
        win.configure(fg_color=COLORS["bg_primary"]); win.attributes("-topmost", True)
        win.grab_set()
        ctk.CTkLabel(win, text="📊 Insert Table", font=FONTS["subheading"],
            text_color=COLORS["text_primary"]).pack(padx=15, pady=(15, 10))
        sf = ctk.CTkFrame(win, fg_color="transparent"); sf.pack(padx=15, pady=5)
        ctk.CTkLabel(sf, text="Columns:", font=FONTS["body"],
            text_color=COLORS["text_secondary"]).pack(side="left", padx=(0, 5))
        cols_var = ctk.StringVar(value="3")
        ctk.CTkEntry(sf, textvariable=cols_var, width=50, height=28,
            fg_color=COLORS["bg_input"], text_color=COLORS["text_primary"],
            border_color=COLORS["border"], corner_radius=6).pack(side="left", padx=(0, 15))
        ctk.CTkLabel(sf, text="Rows:", font=FONTS["body"],
            text_color=COLORS["text_secondary"]).pack(side="left", padx=(0, 5))
        rows_var = ctk.StringVar(value="3")
        ctk.CTkEntry(sf, textvariable=rows_var, width=50, height=28,
            fg_color=COLORS["bg_input"], text_color=COLORS["text_primary"],
            border_color=COLORS["border"], corner_radius=6).pack(side="left")
        def insert():
            try:
                cols = max(1, min(10, int(cols_var.get())))
                rows = max(1, min(20, int(rows_var.get())))
            except ValueError: cols, rows = 3, 3
            header = "| " + " | ".join(f"Header {c+1}" for c in range(cols)) + " |"
            sep = "| " + " | ".join("---" for _ in range(cols)) + " |"
            body = "\n".join("| " + " | ".join("   " for _ in range(cols)) + " |" for _ in range(rows))
            cursor = self.editor.index("insert")
            self.editor.insert(cursor, f"\n{header}\n{sep}\n{body}\n")
            self.editor.focus_set(); win.destroy()
        ctk.CTkButton(win, text="Insert Table", height=34, font=FONTS["body_bold"],
            fg_color=COLORS["accent"], hover_color=COLORS["accent_hover"],
            corner_radius=8, command=insert).pack(padx=15, pady=(15, 10))

    # ── Document Stats ────────────────────────────────────────────

    def _update_stats(self):
        try:
            content = self.editor.get("1.0", "end-1c")
            chars = len(content)
            words = len(content.split()) if content.strip() else 0
            lines = content.count("\n") + 1 if content.strip() else 0
            self.stats_label.configure(
                text=f"Words: {words}  |  Characters: {chars}  |  Lines: {lines}")
        except Exception: pass

    # ── Font Size ─────────────────────────────────────────────────

    def _change_font_size(self, size_str):
        try:
            self.editor.configure(font=("Consolas", int(size_str)))
        except (ValueError, Exception): pass

    # ── Find & Replace ────────────────────────────────────────────

    def _show_find_replace(self):
        win = ctk.CTkToplevel(self)
        win.title("Find & Replace"); win.geometry("420x200")
        win.configure(fg_color=COLORS["bg_primary"]); win.attributes("-topmost", True)
        ctk.CTkLabel(win, text="🔎 Find & Replace", font=FONTS["subheading"],
            text_color=COLORS["text_primary"]).pack(padx=15, pady=(15, 10))
        ff = ctk.CTkFrame(win, fg_color="transparent"); ff.pack(fill="x", padx=15, pady=2)
        ctk.CTkLabel(ff, text="Find:", font=FONTS["body"],
            text_color=COLORS["text_secondary"], width=70).pack(side="left")
        find_e = ctk.CTkEntry(ff, fg_color=COLORS["bg_input"], text_color=COLORS["text_primary"],
            font=FONTS["body"], border_color=COLORS["border"], corner_radius=6)
        find_e.pack(side="left", fill="x", expand=True)
        rf = ctk.CTkFrame(win, fg_color="transparent"); rf.pack(fill="x", padx=15, pady=2)
        ctk.CTkLabel(rf, text="Replace:", font=FONTS["body"],
            text_color=COLORS["text_secondary"], width=70).pack(side="left")
        repl_e = ctk.CTkEntry(rf, fg_color=COLORS["bg_input"], text_color=COLORS["text_primary"],
            font=FONTS["body"], border_color=COLORS["border"], corner_radius=6)
        repl_e.pack(side="left", fill="x", expand=True)
        sl = ctk.CTkLabel(win, text="", font=FONTS["small"], text_color=COLORS["text_muted"])
        sl.pack(padx=15, pady=(2, 0))
        br = ctk.CTkFrame(win, fg_color="transparent"); br.pack(padx=15, pady=(5, 10))
        def find_next():
            q = find_e.get()
            if not q: return
            content = self.editor.get("1.0", "end-1c")
            cursor = self.editor.index("insert")
            start_idx = len(self.editor.get("1.0", cursor))
            pos = content.find(q, start_idx)
            if pos == -1: pos = content.find(q)
            if pos == -1:
                sl.configure(text="Not found", text_color=COLORS["danger"]); return
            line = content[:pos].count("\n") + 1
            col = pos - content[:pos].rfind("\n") - 1
            start = f"{line}.{col}"
            end_pos = pos + len(q)
            end_line = content[:end_pos].count("\n") + 1
            end_col = end_pos - content[:end_pos].rfind("\n") - 1
            end = f"{end_line}.{end_col}"
            self.editor.tag_remove("sel", "1.0", "end")
            self.editor.tag_add("sel", start, end)
            self.editor.mark_set("insert", end); self.editor.see(start)
            sl.configure(text=f"{content.count(q)} match(es) found", text_color=COLORS["success"])
        def replace_one():
            q, r = find_e.get(), repl_e.get()
            if not q: return
            try:
                ss = self.editor.index("sel.first"); se = self.editor.index("sel.last")
                if self.editor.get(ss, se) == q:
                    self.editor.delete(ss, se); self.editor.insert(ss, r)
                    sl.configure(text="Replaced 1 occurrence", text_color=COLORS["success"])
                    self._update_stats(); find_next(); return
            except Exception: pass
            find_next()
        def replace_all():
            q, r = find_e.get(), repl_e.get()
            if not q: return
            content = self.editor.get("1.0", "end-1c")
            count = content.count(q)
            if count == 0:
                sl.configure(text="Not found", text_color=COLORS["danger"]); return
            self.editor.delete("1.0", "end"); self.editor.insert("1.0", content.replace(q, r))
            sl.configure(text=f"Replaced {count} occurrence(s)", text_color=COLORS["success"])
            self._update_stats()
        ctk.CTkButton(br, text="Find", width=80, height=30, font=FONTS["small"],
            fg_color=COLORS["accent"], hover_color=COLORS["accent_hover"],
            corner_radius=6, command=find_next).pack(side="left", padx=3)
        ctk.CTkButton(br, text="Replace", width=80, height=30, font=FONTS["small"],
            fg_color=COLORS["warning"], corner_radius=6, command=replace_one).pack(side="left", padx=3)
        ctk.CTkButton(br, text="Replace All", width=90, height=30, font=FONTS["small"],
            fg_color=COLORS["danger"], corner_radius=6, command=replace_all).pack(side="left", padx=3)

    # ── Export ────────────────────────────────────────────────────

    def _export_note(self):
        if not self.sel_id: return
        note = db.get_note(self.sel_id)
        if not note: return
        try:
            content = self.editor.get("1.0", "end-1c")
            title = self.title_e.get().strip() or note["title"]
        except Exception:
            content = note["content"]; title = note["title"]
        filepath = filedialog.asksaveasfilename(title="Export Note", initialfile=title,
            defaultextension=".md", filetypes=[("Markdown","*.md"),("Text file","*.txt"),
                ("Word Document","*.docx")])
        if not filepath: return
        ext = os.path.splitext(filepath)[1].lower()
        try:
            if ext == ".docx":
                self._export_docx(filepath, title, content)
            else:
                with open(filepath, "w", encoding="utf-8") as f: f.write(content)
            self._show_export_status(f"✅ Exported to {os.path.basename(filepath)}")
        except Exception as e:
            self._show_export_status(f"❌ Export failed: {str(e)[:60]}")

    def _export_docx(self, filepath, title, content):
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            with open(filepath, "w", encoding="utf-8") as f: f.write(content)
            return
        doc = Document()
        doc.add_heading(title, level=0)
        for line in content.split("\n"):
            stripped = line.strip()
            if stripped.startswith("#### "): doc.add_heading(stripped[5:], level=4)
            elif stripped.startswith("### "): doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "): doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "): doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("---") or stripped.startswith("***"):
                p = doc.add_paragraph(); p.add_run("─" * 50)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif stripped.startswith("- [ ] "):
                doc.add_paragraph(stripped[6:], style="List Bullet")
            elif stripped.startswith("- [x] ") or stripped.startswith("- [X] "):
                doc.add_paragraph(f"✓ {stripped[6:]}", style="List Bullet")
            elif stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif re.match(r'^\d+\.\s', stripped):
                doc.add_paragraph(re.sub(r'^\d+\.\s', '', stripped), style="List Number")
            elif stripped.startswith("> "):
                p = doc.add_paragraph(); run = p.add_run(stripped[2:]); run.italic = True
            elif stripped.startswith("```"): continue
            elif stripped == "": doc.add_paragraph()
            else:
                p = doc.add_paragraph(); self._docx_format_line(p, stripped)
        doc.save(filepath)

    def _docx_format_line(self, paragraph, text):
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|~~.*?~~)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2]); run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1]); run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1]); run.font.name = "Consolas"
            elif part.startswith("~~") and part.endswith("~~"):
                run = paragraph.add_run(part[2:-2]); run.font.strike = True
            elif part: paragraph.add_run(part)

    def _show_export_status(self, message):
        try:
            self.stats_label.configure(text=message)
            self.after(3000, self._update_stats)
        except Exception: pass

    # ── Preview ───────────────────────────────────────────────────

    def _toggle_preview(self):
        self.preview_visible = not self.preview_visible
        try:
            self.preview_btn.configure(
                fg_color=COLORS["accent"] if self.preview_visible else COLORS["bg_card"])
        except Exception:
            pass
        self._update_editor_layout()

    def _render_preview(self):
        content = self.editor.get("1.0", "end").strip()
        self.preview_panel.configure(state="normal")
        self.preview_panel.delete("1.0", "end")
        if not content:
            self.preview_panel.insert("1.0", "Start typing to see preview...")
            self.preview_panel.configure(state="disabled")
            return
        rendered = self._markdown_to_display(content)
        self.preview_panel.insert("1.0", rendered)
        self.preview_panel.configure(state="disabled")

    def _markdown_to_display(self, text):
        lines = text.split("\n")
        output = []
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("#### "):
                output.append(f"    {stripped[5:]}")
            elif stripped.startswith("### "):
                output.append(f"   {stripped[4:].upper()}")
            elif stripped.startswith("## "):
                output.append(f"  ━━ {stripped[3:].upper()} ━━")
            elif stripped.startswith("# "):
                output.append(f"═══ {stripped[2:].upper()} ═══")
            elif stripped.startswith("---") or stripped.startswith("***"):
                output.append("─" * 40)
            elif stripped.startswith("- [ ] "):
                output.append(f"  ☐ {stripped[6:]}")
            elif stripped.startswith("- [x] ") or stripped.startswith("- [X] "):
                output.append(f"  ☑ {stripped[6:]}")
            elif stripped.startswith("- "):
                output.append(f"  • {stripped[2:]}")
            elif stripped.startswith("> "):
                output.append(f"  │ {stripped[2:]}")
            elif stripped.startswith("```"):
                output.append("  ┌────────────────────")
            else:
                display = stripped
                display = re.sub(r'\*\*(.+?)\*\*', r'[\1]', display)
                display = re.sub(r'\*(.+?)\*', r'\1', display)
                display = re.sub(r'~~(.+?)~~', r'\1', display)
                display = re.sub(r'`(.+?)`', r'⟨\1⟩', display)
                output.append(display)
        return "\n".join(output)

    # ── Markdown Navigator ────────────────────────────────────────

    def _toggle_navigator(self):
        self.nav_visible = not self.nav_visible
        try:
            self.nav_btn.configure(
                fg_color=COLORS["accent"] if self.nav_visible else COLORS["bg_card"])
        except Exception:
            pass
        self._update_editor_layout()

    def _refresh_navigator(self):
        for w in self.nav_panel.winfo_children(): w.destroy()
        ctk.CTkLabel(self.nav_panel, text="Headings", font=FONTS["body_bold"],
            text_color=COLORS["text_primary"]).pack(anchor="w", padx=6, pady=(4, 6))
        content = self.editor.get("1.0", "end")
        lines = content.split("\n")
        found = False
        for i, line in enumerate(lines):
            stripped = line.strip()
            level, title = 0, ""
            if stripped.startswith("#### "):
                level, title = 4, stripped[5:]
            elif stripped.startswith("### "):
                level, title = 3, stripped[4:]
            elif stripped.startswith("## "):
                level, title = 2, stripped[3:]
            elif stripped.startswith("# "):
                level, title = 1, stripped[2:]
            if level and title:
                found = True
                indent = (level - 1) * 10
                line_num = i + 1
                ctk.CTkButton(self.nav_panel, text=title[:30],
                    font=FONTS["small"] if level > 2 else FONTS["body_bold"],
                    fg_color="transparent", hover_color=COLORS["bg_secondary"],
                    text_color=COLORS["text_secondary"], anchor="w", height=24, corner_radius=4,
                    command=lambda ln=line_num: self._navigate_to_line(ln)
                ).pack(fill="x", padx=(6 + indent, 4), pady=1)
        if not found:
            ctk.CTkLabel(self.nav_panel, text="No headings found", font=FONTS["small"],
                text_color=COLORS["text_muted"]).pack(padx=6, pady=8)

    def _navigate_to_line(self, line_num):
        pos = f"{line_num}.0"
        self.editor.mark_set("insert", pos)
        self.editor.see(pos)
        self.editor.focus_set()

    # ── Focus Mode ────────────────────────────────────────────────

    def _toggle_focus_mode(self):
        self.app.toggle_focus_mode()

    def update_focus_btn(self, is_focused):
        if is_focused:
            self.focus_btn.configure(text="↩️ Exit Focus", fg_color=COLORS["accent"])
        else:
            self.focus_btn.configure(text="🖥️ Focus Mode", fg_color=COLORS["bg_card"])

    # ── Note CRUD ─────────────────────────────────────────────────

    def _new_note(self):
        """Create a new blank note and open it in the editor."""
        note_id = db.add_note("Untitled Note", "")
        if not note_id:
            return
        self.refresh()
        self.view_note(note_id)

    def _save(self, nid):
        db.update_note(nid, title=self.title_e.get().strip(),
            content=self.editor.get("1.0","end").strip(), tags=self.tags_e.get().strip())
        self.refresh()

    def _del(self, nid):
        db.delete_note(nid); self.sel_id = None; self.refresh(); self._empty_viewer()

    def _gen_cards(self, note):
        self.app.select_tab("Flashcards")
        fc = self.app.tabs.get("Flashcards")
        if fc: fc.show_ai_gen()

    def _summarize(self, note):
        if not self.app.claude_client: return
        def run():
            try:
                r = self.app.claude_client.summarize_notes(note["content"])
                self.after(0, lambda: self._popup("📋 Summary", r))
            except Exception as e:
                self.after(0, lambda: self._popup("Error", str(e)))
        threading.Thread(target=run, daemon=True).start()

    def _ask(self, note):
        d = ctk.CTkInputDialog(text="Ask a question about this note:", title="Ask Claude")
        q = d.get_input()
        if not q or not self.app.claude_client: return
        def run():
            try:
                r = self.app.claude_client.answer_question(q, note["content"])
                self.after(0, lambda: self._popup(f"❓ {q}", r))
            except Exception as e:
                self.after(0, lambda: self._popup("Error", str(e)))
        threading.Thread(target=run, daemon=True).start()

    def _popup(self, title, content):
        w = ctk.CTkToplevel(self); w.title("StudyForge"); w.geometry("700x500")
        w.configure(fg_color=COLORS["bg_primary"]); w.attributes("-topmost", True)
        ctk.CTkLabel(w, text=title, font=FONTS["subheading"],
            text_color=COLORS["accent_light"]).pack(padx=15, pady=(15,5), anchor="w")
        tb = ctk.CTkTextbox(w, fg_color=COLORS["bg_input"], text_color=COLORS["text_primary"],
            font=FONTS["body"], wrap="word", corner_radius=8)
        tb.insert("1.0", content); tb.configure(state="disabled")
        tb.pack(fill="both", expand=True, padx=15, pady=(0,15))

    def import_file(self):
        fps = filedialog.askopenfilenames(title="Import Notes",
            filetypes=[("All Supported","*.txt *.md *.pdf *.docx"),("Text","*.txt"),
                       ("Markdown","*.md"),("PDF","*.pdf"),("Word","*.docx")])
        for fp in fps:
            fn = os.path.basename(fp)
            db.add_note(os.path.splitext(fn)[0], extract_text(fp), source_file=fn)
        if fps: self.refresh()

    def paste_dlg(self):
        w = ctk.CTkToplevel(self); w.title("Paste Note"); w.geometry("600x450")
        w.configure(fg_color=COLORS["bg_primary"]); w.attributes("-topmost", True)
        ctk.CTkLabel(w, text="Title:", font=FONTS["body"],
            text_color=COLORS["text_secondary"]).pack(padx=15, pady=(15,2), anchor="w")
        te = ctk.CTkEntry(w, fg_color=COLORS["bg_input"], text_color=COLORS["text_primary"],
            font=FONTS["body"], corner_radius=8)
        te.pack(fill="x", padx=15)
        ctk.CTkLabel(w, text="Content:", font=FONTS["body"],
            text_color=COLORS["text_secondary"]).pack(padx=15, pady=(10,2), anchor="w")
        cb = ctk.CTkTextbox(w, fg_color=COLORS["bg_input"], text_color=COLORS["text_primary"],
            font=FONTS["body"], wrap="word", corner_radius=8)
        cb.pack(fill="both", expand=True, padx=15, pady=(0,10))
        def save():
            t = te.get().strip() or "Untitled"
            c = cb.get("1.0","end").strip()
            if c: db.add_note(t, c); self.refresh(); w.destroy()
        ctk.CTkButton(w, text="💾 Save", height=36, font=FONTS["body_bold"],
            fg_color=COLORS["success"], corner_radius=8, command=save).pack(padx=15, pady=(0,15))
